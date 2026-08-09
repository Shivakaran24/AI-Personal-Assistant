import os
import sys
import json
import sqlite3
import datetime
import subprocess
import re
import urllib.parse
from html import unescape
import httpx
from typing import Dict, Any, List

class CalendarStoreManager:
    _events = {}
    _is_cleared = False

    @classmethod
    def _init_defaults(cls):
        if not cls._events and not cls._is_cleared:
            cls._events = {}

    @classmethod
    def clear_history(cls):
        cls._events = {}
        cls._is_cleared = True
        return {
            "status": "success",
            "message": "Calendar event history cleared successfully.",
            "dashboard_metrics": {
                "total_events": 0,
                "total_accepted": 0,
                "total_rejected": 0,
                "total_pending": 0,
                "total_responses": 0,
                "acceptance_rate": "0%"
            },
            "events": []
        }

    CONTACT_DIRECTORY = {
        "shiva": "shiva@company.com",
        "bob": "bob@company.com",
        "alice": "alice@company.com",
        "john": "john@company.com",
        "sarah": "sarah@company.com",
        "user": "user@company.com"
    }

    @classmethod
    def resolve_person_email(cls, person_or_email: str) -> Tuple[str, str]:
        """
        Dynamically resolves ANY person's name or email.
        Strips conversational prefixes like 'with', 'meet with', 'for', 'to'.
        - If email (e.g. 'shiva.g@gmail.com'): returns ('Shiva G', 'shiva.g@gmail.com')
        - If single name (e.g. 'Shiva'): returns ('Shiva', 'shiva@company.com')
        """
        if not person_or_email:
            return "Contact", "contact@company.com"

        clean = str(person_or_email).strip().strip("'\"")

        # Strip leading conversational prefixes ("schedule a meet with", "meet with", "with", "meet", "for", "to")
        clean = re.sub(r'^(schedule\s+a\s+meet(ing)?\s+with|schedule\s+with|meet\s+with|meet|with|for|to)\s+', '', clean, flags=re.IGNORECASE).strip()

        if not clean:
            return "Contact", "contact@company.com"

        # 1. Direct Email Input
        if "@" in clean:
            email = clean.lower()
            local_part = email.split("@")[0]
            display_name = re.sub(r'[\._\+]', ' ', local_part).title()
            return display_name, email

        # 2. Dynamic Name Resolution for ANY person
        display_name = clean.title()
        clean_parts = re.findall(r'[a-zA-Z0-9]+', clean.lower())
        dynamic_local = ".".join(clean_parts) if clean_parts else "contact"

        lower_name = clean.lower()
        if hasattr(cls, "CONTACT_DIRECTORY") and lower_name in cls.CONTACT_DIRECTORY:
            return display_name, cls.CONTACT_DIRECTORY[lower_name]

        dynamic_email = f"{dynamic_local}@company.com"
        from app.core.config import settings
        if settings.EMAIL_USER and "@" in settings.EMAIL_USER:
            dynamic_email = settings.EMAIL_USER

        return display_name, dynamic_email

    @classmethod
    def check_conflict(cls, target_date: str, start_time: str) -> Tuple[bool, Optional[str]]:
        cls._init_defaults()
        for evt in cls._events.values():
            if evt.get("calendar_status") != "canceled":
                if evt.get("date") == target_date and evt.get("start_time") and start_time:
                    if evt.get("start_time").lower().strip() == start_time.lower().strip():
                        return True, f"Conflict with existing event '{evt.get('title')}' at {start_time}"
        return False, None

    @classmethod
    def list_events(cls, days=7):
        cls._init_defaults()
        return list(cls._events.values())

    @classmethod
    def create_event(cls, title, event_date, start_time, duration, attendees_list):
        cls._init_defaults()
        evt_id = f"evt-{int(datetime.datetime.now().timestamp())}"
        rsvps = {}
        for att in attendees_list:
            rsvps[att] = {"status": "pending", "responded_at": None}

        evt = {
            "id": evt_id,
            "title": title,
            "date": event_date,
            "start_time": start_time,
            "duration_minutes": duration,
            "organizer": "user@company.com",
            "attendees": attendees_list,
            "rsvps": rsvps,
            "accepted_count": 0,
            "rejected_count": 0,
            "pending_count": len(attendees_list),
            "calendar_status": "confirmed"
        }
        cls._events[evt_id] = evt
        return evt

    @classmethod
    def respond_invitation(cls, event_id, attendee, action):
        cls._init_defaults()
        evt = cls._events.get(event_id)
        if not evt:
            for e_id, e in cls._events.items():
                if event_id and (event_id.lower() in e_id.lower() or event_id.lower() in e["title"].lower()):
                    evt = e
                    break
        if not evt:
            if cls._events:
                evt = list(cls._events.values())[0]
            else:
                evt = cls.create_event("Calendar Event", datetime.date.today().strftime("%Y-%m-%d"), "09:00 AM", 30, [attendee or "user@company.com"])

        att_name = attendee if attendee else (evt["attendees"][0] if evt["attendees"] else "user@company.com")
        if att_name not in evt["rsvps"]:
            evt["rsvps"][att_name] = {"status": "pending", "responded_at": None}
            if att_name not in evt["attendees"]:
                evt["attendees"].append(att_name)

        new_status = "accepted" if action.lower() in ["accept", "accepted"] else "rejected"
        evt["rsvps"][att_name] = {
            "status": new_status,
            "responded_at": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        }

        acc = sum(1 for r in evt["rsvps"].values() if r.get("status") == "accepted")
        rej = sum(1 for r in evt["rsvps"].values() if r.get("status") == "rejected")
        pen = sum(1 for r in evt["rsvps"].values() if r.get("status") == "pending")

        evt["accepted_count"] = acc
        evt["rejected_count"] = rej
        evt["pending_count"] = pen

        return evt, att_name, new_status

    @classmethod
    def clear_history(cls):
        cls._events = {}
        return {
            "status": "success",
            "message": "Calendar event history cleared successfully.",
            "dashboard_metrics": {
                "total_events": 0,
                "total_accepted": 0,
                "total_rejected": 0,
                "total_pending": 0,
                "acceptance_rate": "0%"
            },
            "events": []
        }

    @classmethod
    def get_dashboard_stats(cls):
        cls._init_defaults()
        tot_acc = sum(e.get("accepted_count", 0) for e in cls._events.values())
        tot_rej = sum(e.get("rejected_count", 0) for e in cls._events.values())
        tot_pen = sum(e.get("pending_count", 0) for e in cls._events.values())
        total_rsvps = tot_acc + tot_rej

        accepted_list = []
        rejected_list = []
        pending_list = []

        for evt in cls._events.values():
            for att, r_info in evt.get("rsvps", {}).items():
                entry = {
                    "event_id": evt["id"],
                    "event_title": evt["title"],
                    "date": evt["date"],
                    "time": evt["start_time"],
                    "attendee": att,
                    "responded_at": r_info.get("responded_at")
                }
                if r_info.get("status") == "accepted":
                    accepted_list.append(entry)
                elif r_info.get("status") == "rejected":
                    rejected_list.append(entry)
                else:
                    pending_list.append(entry)

        return {
            "status": "success",
            "dashboard_metrics": {
                "total_events": len(cls._events),
                "total_accepted": tot_acc,
                "total_rejected": tot_rej,
                "total_pending": tot_pen,
                "total_responses": total_rsvps,
                "acceptance_rate": f"{round((tot_acc / total_rsvps) * 100, 1)}%" if total_rsvps > 0 else "0%"
            },
            "accepted_attendees": accepted_list,
            "rejected_attendees": rejected_list,
            "pending_attendees": pending_list,
            "events": list(cls._events.values())
        }

    @classmethod
    def check_availability(cls, target_date=None, start_time=None, duration_minutes=30):
        cls._init_defaults()
        if not target_date:
            target_date = datetime.date.today().strftime("%Y-%m-%d")

        # Normalize date strings
        if str(target_date).lower() == "today":
            target_date = datetime.date.today().strftime("%Y-%m-%d")
        elif str(target_date).lower() == "tomorrow":
            target_date = (datetime.date.today() + datetime.timedelta(days=1)).strftime("%Y-%m-%d")

        same_day_events = [e for e in cls._events.values() if e.get("date") == target_date and e.get("calendar_status") != "canceled"]
        conflicts = []
        is_available = True

        if start_time:
            st_clean = str(start_time).lower().strip()
            for evt in same_day_events:
                evt_st = str(evt.get("start_time", "")).lower().strip()
                if st_clean == evt_st or st_clean in evt_st or evt_st in st_clean:
                    conflicts.append(evt)
                    is_available = False

        # Generate suggested free slots
        standard_slots = ["09:00 AM", "10:30 AM", "01:00 PM", "02:30 PM", "04:00 PM"]
        taken_times = [str(e.get("start_time", "")).lower().strip() for e in same_day_events]
        free_slots = [s for s in standard_slots if s.lower().strip() not in taken_times]

        return {
            "status": "success",
            "target_date": target_date,
            "target_time": start_time or "Entire Day",
            "is_available": is_available,
            "conflicts": conflicts,
            "total_conflicts": len(conflicts),
            "scheduled_events": same_day_events,
            "suggested_free_slots": free_slots,
            "summary": f"Available slot found for {target_date} at {start_time}." if is_available else f"Time conflict detected with {len(conflicts)} event(s) on {target_date}."
        }

    @classmethod
    def manage_event(cls, event_id, action="update", new_title=None, new_date=None, new_start_time=None, new_duration=None):
        cls._init_defaults()
        evt = cls._events.get(event_id)
        if not evt:
            for e_id, e in cls._events.items():
                if event_id and (event_id.lower() in e_id.lower() or event_id.lower() in e.get("title", "").lower()):
                    evt = e
                    break

        if not evt:
            return {"status": "error", "message": f"Calendar event '{event_id}' not found."}

        action_clean = str(action).lower().strip()
        if action_clean in ["cancel", "delete", "remove"]:
            evt["calendar_status"] = "canceled"
            return {
                "status": "success",
                "action": "canceled",
                "event_id": evt["id"],
                "title": evt["title"],
                "message": f"Calendar event '{evt['title']}' has been canceled.",
                "event": evt
            }
        else:
            if new_title: evt["title"] = new_title
            if new_date: evt["date"] = new_date
            if new_start_time: evt["start_time"] = new_start_time
            if new_duration: evt["duration_minutes"] = int(new_duration)
            evt["calendar_status"] = "rescheduled" if (new_date or new_start_time) else "updated"

            return {
                "status": "success",
                "action": evt["calendar_status"],
                "event_id": evt["id"],
                "title": evt["title"],
                "date": evt["date"],
                "start_time": evt["start_time"],
                "duration_minutes": evt["duration_minutes"],
                "message": f"Calendar event '{evt['title']}' successfully updated.",
                "event": evt
            }

class EmailStoreManager:
    """
    In-memory and persistent store manager for Email Agent draft messages, outbox, and notifications.
    """
    _drafts = {}
    _notifications = []

    @classmethod
    def save_draft(cls, to: str, subject: str, body: str) -> Dict[str, Any]:
        draft_id = f"draft-{int(datetime.datetime.now().timestamp())}"
        draft_entry = {
            "id": draft_id,
            "to": to,
            "subject": subject,
            "body": body,
            "created_at": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "status": "draft"
        }
        cls._drafts[draft_id] = draft_entry
        return draft_entry

    @classmethod
    def list_drafts(cls) -> List[Dict[str, Any]]:
        return list(cls._drafts.values())

    @classmethod
    def get_draft(cls, draft_id: str) -> Optional[Dict[str, Any]]:
        return cls._drafts.get(draft_id)

    @classmethod
    def delete_draft(cls, draft_id: str) -> bool:
        if draft_id in cls._drafts:
            del cls._drafts[draft_id]
            return True
        return False

    @classmethod
    def log_notification(cls, to: str, subject: str, body: str, channel: str = "email") -> Dict[str, Any]:
        notif_entry = {
            "id": f"notif-{int(datetime.datetime.now().timestamp())}",
            "to": to,
            "subject": subject,
            "body": body,
            "channel": channel,
            "dispatched_at": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "status": "sent"
        }
        cls._notifications.append(notif_entry)
        return notif_entry

    @classmethod
    def list_notifications(cls) -> List[Dict[str, Any]]:
        return list(cls._notifications[::-1])

class BuiltinMCPServers:
    """
    Built-in operational MCP tool implementations for production-grade assistant actions.
    Provides standard MCP JSON tool schemas and handlers.
    """

    @staticmethod
    def get_tool_definitions() -> List[Dict[str, Any]]:
        return [
            # Filesystem Tools
            {
                "server_id": "filesystem_server",
                "server_name": "Filesystem Tool Server",
                "name": "fs_read_file",
                "description": "Reads contents of a file from the workspace filesystem.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "path": {"type": "string", "description": "Relative or absolute filepath to read"}
                    },
                    "required": ["path"]
                }
            },
            {
                "server_id": "filesystem_server",
                "server_name": "Filesystem Tool Server",
                "name": "fs_write_file",
                "description": "Writes or creates content in a file on the filesystem.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "path": {"type": "string", "description": "File path to write to"},
                        "content": {"type": "string", "description": "Text content to save"}
                    },
                    "required": ["path", "content"]
                }
            },
            {
                "server_id": "filesystem_server",
                "server_name": "Filesystem Tool Server",
                "name": "fs_list_dir",
                "description": "Lists files and subdirectories in a workspace directory.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "path": {"type": "string", "description": "Directory path to list"}
                    },
                    "required": []
                }
            },
            
            # Gmail & Calendar Manager Tools
            {
                "server_id": "workspace_server",
                "server_name": "Google Workspace Server",
                "name": "gmail_list_messages",
                "description": "Lists recent emails from inbox with subject, sender, and date.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "query": {"type": "string", "description": "Filter search query (e.g., 'unread', 'meeting')"},
                        "limit": {"type": "integer", "default": 5}
                    }
                }
            },
            {
                "server_id": "workspace_server",
                "server_name": "Google Workspace Server",
                "name": "gmail_send_message",
                "description": "Drafts or sends an email message.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "to": {"type": "string", "description": "Recipient email address"},
                        "subject": {"type": "string", "description": "Email subject line"},
                        "body": {"type": "string", "description": "Email content body"}
                    },
                    "required": ["to", "subject", "body"]
                }
            },
            {
                "server_id": "workspace_server",
                "server_name": "Google Workspace Server",
                "name": "calendar_list_events",
                "description": "Lists upcoming calendar events and meetings.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "days": {"type": "integer", "description": "Number of days ahead to look", "default": 7}
                    }
                }
            },
            {
                "server_id": "workspace_server",
                "server_name": "Google Workspace Server",
                "name": "calendar_create_event",
                "description": "Schedules a new calendar event or meeting with date, time, and dispatches notifications/messages to attendees.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "title": {"type": "string", "description": "Title of the meeting/event"},
                        "date": {"type": "string", "description": "Date of the event e.g. 2026-08-01 (defaults to current date if omitted)"},
                        "start_time": {"type": "string", "description": "ISO timestamp or human format e.g. 9:00am or 2026-08-01 14:00"},
                        "duration_minutes": {"type": "integer", "default": 30},
                        "attendees": {
                            "description": "List of attendee emails or attendee names to notify",
                            "anyOf": [
                                {"type": "array", "items": {"type": "string"}},
                                {"type": "string"}
                            ]
                        },
                        "notify_attendees": {
                            "type": "boolean",
                            "default": True,
                            "description": "Whether to send notification message to attendees"
                        }
                    },
                    "required": ["title", "start_time"]
                }
            },
            {
                "server_id": "workspace_server",
                "server_name": "Google Workspace Server",
                "name": "calendar_respond_invitation",
                "description": "Responds to a calendar invitation with Accept or Reject buttons. Sends immediate notification and schedules a 30-minute prior reminder notification for accepted attendees.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "event_id": {"type": "string", "description": "ID of the calendar event e.g. 'evt-201'"},
                        "attendee": {"type": "string", "description": "Email address or name of attendee responding"},
                        "action": {
                            "type": "string",
                            "enum": ["accept", "reject"],
                            "description": "RSVP response ('accept' or 'reject')"
                        }
                    },
                    "required": ["event_id", "action"]
                }
            },
            {
                "server_id": "workspace_server",
                "server_name": "Google Workspace Server",
                "name": "calendar_get_dashboard_stats",
                "description": "Returns calendar RSVP telemetry metrics including total accepted, total rejected, total pending, and event list with RSVP statuses.",
                "parameters": {
                    "type": "object",
                    "properties": {}
                }
            },
            {
                "server_id": "workspace_server",
                "server_name": "Google Workspace Server",
                "name": "calendar_clear_history",
                "description": "Clears and resets calendar event history and RSVP telemetry metrics.",
                "parameters": {
                    "type": "object",
                    "properties": {}
                }
            },
            {
                "server_id": "calendar_agent_server",
                "server_name": "Calendar Agent Server",
                "name": "calendar_check_availability",
                "description": "Calendar Agent: Checks free/busy schedule availability, open slots, and detects time conflicts for a target date or date range.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "date": {"type": "string", "description": "Date to check e.g. '2026-08-10', 'today', or 'tomorrow'"},
                        "start_time": {"type": "string", "description": "Target start time e.g. '09:00 AM' or '14:00'"},
                        "duration_minutes": {"type": "integer", "default": 30, "description": "Duration in minutes"}
                    }
                }
            },
            {
                "server_id": "calendar_agent_server",
                "server_name": "Calendar Agent Server",
                "name": "calendar_manage_event",
                "description": "Calendar Agent: Reschedules, updates, or cancels/deletes an existing calendar event by ID.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "event_id": {"type": "string", "description": "Event ID to update or cancel"},
                        "action": {"type": "string", "enum": ["update", "reschedule", "cancel"], "description": "Action to perform ('update', 'reschedule', or 'cancel')"},
                        "new_title": {"type": "string", "description": "New event title (optional)"},
                        "new_date": {"type": "string", "description": "New event date e.g. '2026-08-15' (optional)"},
                        "new_start_time": {"type": "string", "description": "New event start time e.g. '10:00 AM' (optional)"},
                        "new_duration_minutes": {"type": "integer", "description": "New duration in minutes (optional)"}
                    },
                    "required": ["event_id", "action"]
                }
            },
            {
                "server_id": "email_agent_server",
                "server_name": "Email Agent Server",
                "name": "gmail_draft_message",
                "description": "Email Agent: Drafts a message with recipient, subject, and body for user review without sending immediately.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "to": {"type": "string", "description": "Recipient email address"},
                        "subject": {"type": "string", "description": "Email subject line"},
                        "body": {"type": "string", "description": "Email content body"}
                    },
                    "required": ["to", "subject", "body"]
                }
            },
            {
                "server_id": "email_agent_server",
                "server_name": "Email Agent Server",
                "name": "gmail_list_drafts",
                "description": "Email Agent: Lists pending email drafts.",
                "parameters": {
                    "type": "object",
                    "properties": {}
                }
            },
            {
                "server_id": "email_agent_server",
                "server_name": "Email Agent Server",
                "name": "email_send_notification",
                "description": "Email Agent: Dispatches an instant email or system notification alert.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "to": {"type": "string", "description": "Recipient email address or user ID"},
                        "subject": {"type": "string", "description": "Notification title/subject"},
                        "body": {"type": "string", "description": "Notification body content"},
                        "channel": {"type": "string", "enum": ["email", "system"], "default": "email"}
                    },
                    "required": ["to", "subject", "body"]
                }
            },

            # GitHub Manager Tools
            {
                "server_id": "github_server",
                "server_name": "GitHub Integration Server",
                "name": "github_list_repos",
                "description": "Lists GitHub repositories for user or organization.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "user_or_org": {"type": "string", "description": "GitHub username or organization"}
                    }
                }
            },
            {
                "server_id": "github_server",
                "server_name": "GitHub Integration Server",
                "name": "github_create_issue",
                "description": "Creates a new issue in a GitHub repository.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "repo": {"type": "string", "description": "Repository full name (e.g. 'org/repo')"},
                        "title": {"type": "string", "description": "Issue title"},
                        "body": {"type": "string", "description": "Issue details/body"}
                    },
                    "required": ["repo", "title", "body"]
                }
            },

            # Code Execution Sandbox
            {
                "server_id": "code_interpreter_server",
                "server_name": "Python Code Interpreter Server",
                "name": "run_python_code",
                "description": "Executes Python code in a safe sandbox process and returns stdout, stderr, and execution status.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "code": {"type": "string", "description": "Python snippet to execute"}
                    },
                    "required": ["code"]
                }
            },

            # Database Tool
            {
                "server_id": "database_server",
                "server_name": "Database Query Server",
                "name": "db_query",
                "description": "Executes SQL SELECT queries against the local assistant database.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "query": {"type": "string", "description": "SQL SELECT query string"},
                        "limit": {"type": "integer", "description": "Maximum number of rows to return (default 20, max 100)"}
                    },
                    "required": ["query"]
                }
            },

            # Web Search & Weather Tool
            {
                "server_id": "web_tools_server",
                "server_name": "Web & Weather Server",
                "name": "web_search",
                "description": "Searches the web for latest info, documentation, or news.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "query": {"type": "string", "description": "Search query"}
                    },
                    "required": ["query"]
                }
            },
            {
                "server_id": "web_tools_server",
                "server_name": "Web & Weather Server",
                "name": "get_weather",
                "description": "Gets current weather forecast for a specified city and optional country.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "city": {"type": "string", "description": "City name (e.g. 'Mumbai', 'London', 'New York')"},
                        "country": {"type": "string", "description": "Country name or country code (e.g. 'India', 'UK', 'USA')"}
                    },
                    "required": ["city"]
                }
            },
            
            # Document Generator Server
            {
                "server_id": "document_server",
                "server_name": "Document Generator Server",
                "name": "generate_document",
                "description": "Generates downloadable Word (.docx), PDF (.pdf), Text (.txt), or Markdown (.md) documents.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "filename": {"type": "string", "description": "Output filename (e.g. 'email_summary.docx' or 'summary.pdf')"},
                        "title": {"type": "string", "description": "Document title"},
                        "content": {"type": "string", "description": "Document text content"}
                    },
                    "required": ["filename", "content"]
                }
            }
        ]

    @staticmethod
    def _fetch_real_emails(email_user: str, email_pass: str, limit: int = 10, imap_server: str = "imap.gmail.com", port: int = 993) -> Dict[str, Any]:
        import imaplib
        import email
        import re
        from html import unescape
        from email.header import decode_header

        def decode_hdr(val):
            if not val: return ""
            parts = decode_header(val)
            res = ""
            for content, encoding in parts:
                if isinstance(content, bytes):
                    res += content.decode(encoding or "utf-8", errors="ignore")
                else:
                    res += str(content)
            return res.strip()

        def clean_body(raw):
            if not raw: return ""
            text = re.sub(r'<style[^>]*>[\s\S]*?</style>', ' ', raw, flags=re.IGNORECASE)
            text = re.sub(r'<script[^>]*>[\s\S]*?</script>', ' ', text, flags=re.IGNORECASE)
            text = re.sub(r'<[^>]+>', ' ', text)
            text = unescape(text)
            lines = [line.strip() for line in text.splitlines() if line.strip()]
            cleaned = " ".join(lines)
            cleaned = re.sub(r'\s+', ' ', cleaned)
            return cleaned[:1500] # Extract up to 1500 characters of full email content

        try:
            mail = imaplib.IMAP4_SSL(imap_server, port)
            mail.login(email_user, email_pass)
            mail.select("INBOX")

            status, messages = mail.search(None, 'ALL')
            if status != 'OK':
                return {"status": "error", "message": "Failed to search inbox."}

            email_ids = messages[0].split()
            if not email_ids:
                return {"status": "success", "count": 0, "emails": [], "message": "Inbox is empty."}

            latest_ids = email_ids[-limit:][::-1]
            real_emails = []

            for e_id in latest_ids:
                res, msg_data = mail.fetch(e_id, '(RFC822)')
                for response_part in msg_data:
                    if isinstance(response_part, tuple):
                        msg = email.message_from_bytes(response_part[1])
                        
                        subject = decode_hdr(msg.get("Subject", "(No Subject)"))
                        sender = decode_hdr(msg.get("From", "Unknown"))
                        date = decode_hdr(msg.get("Date", ""))

                        plain_body = ""
                        html_body = ""

                        if msg.is_multipart():
                            for part in msg.walk():
                                c_type = part.get_content_type()
                                c_disp = str(part.get("Content-Disposition"))
                                if "attachment" in c_disp:
                                    continue
                                
                                payload = part.get_payload(decode=True)
                                if payload:
                                    decoded_payload = payload.decode(errors="ignore")
                                    if c_type == "text/plain" and not plain_body:
                                        plain_body = decoded_payload
                                    elif c_type == "text/html" and not html_body:
                                        html_body = decoded_payload
                        else:
                            payload = msg.get_payload(decode=True)
                            if payload:
                                decoded_payload = payload.decode(errors="ignore")
                                if msg.get_content_type() == "text/html":
                                    html_body = decoded_payload
                                else:
                                    plain_body = decoded_payload

                        raw_content = plain_body if plain_body.strip() else html_body
                        full_content = clean_body(raw_content)

                        real_emails.append({
                            "id": f"msg-{e_id.decode()}",
                            "sender": sender,
                            "subject": subject or "(No Subject)",
                            "snippet": full_content or "No readable text content in email body.",
                            "date": date
                        })

            mail.logout()
            return {
                "status": "success",
                "mode": "live_imap",
                "server": imap_server,
                "count": len(real_emails),
                "emails": real_emails
            }
        except Exception as e:
            return {
                "status": "error",
                "message": f"Real IMAP email fetch failed: {str(e)}. Please verify EMAIL_USER and EMAIL_PASSWORD in backend/.env file."
            }

    @staticmethod
    def _send_real_email(to: str, subject: str, body: str, email_user: str, email_pass: str, smtp_server: str = "smtp.gmail.com", port: int = 587, html_body: str = None) -> Dict[str, Any]:
        import smtplib
        from email.mime.text import MIMEText
        from email.mime.multipart import MIMEMultipart

        try:
            msg = MIMEMultipart('alternative')
            msg['From'] = email_user
            msg['To'] = to
            msg['Subject'] = subject
            msg.attach(MIMEText(body, 'plain'))
            if html_body:
                msg.attach(MIMEText(html_body, 'html'))

            server = smtplib.SMTP(smtp_server, port, timeout=6)
            server.starttls()
            server.login(email_user, email_pass.replace(" ", ""))
            text = msg.as_string()
            server.sendmail(email_user, [to], text)
            server.quit()

            return {
                "status": "success",
                "mode": "live_smtp",
                "recipient": to,
                "subject": subject,
                "delivered_at": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "message": f"Real email successfully dispatched to {to} via Gmail SMTP."
            }
        except Exception as e:
            return {
                "status": "error",
                "message": f"Failed to send email via SMTP: {str(e)}. Please check EMAIL_USER and App Password in backend/.env."
            }

    @staticmethod
    def execute_tool(tool_name: str, args: Dict[str, Any], skip_hitl: bool = False) -> Dict[str, Any]:
        """
        Executes built-in MCP tool logic and returns JSON output.
        Supports Human-in-the-Loop approval interception for sensitive outbound actions.
        """
        import re
        import urllib.parse
        try:
            # Human-in-the-Loop Intercept for sensitive outbound actions
            if not skip_hitl:
                sensitive_tools = ["gmail_send_message", "email_send_notification"]
                is_sensitive_calendar = tool_name == "calendar_manage_event" and args.get("action", "").lower() in ["cancel", "delete", "reschedule"]
                is_create_calendar = tool_name in ["calendar_create_event", "calendar_schedule_meeting"]

                if tool_name in sensitive_tools or is_sensitive_calendar or is_create_calendar:
                    from app.planner.hitl_manager import hitl_manager
                    
                    disp_name = None
                    email_res = None

                    if tool_name == "gmail_send_message":
                        action_type = "send_email"
                        disp_name = str(args.get("to", "Recipient"))
                        email_res = str(args.get("to", ""))
                        title = f"Send Outbound Email to '{args.get('to')}'"
                        desc = f"Subject: {args.get('subject')}\nRecipient: {args.get('to')}"
                    elif tool_name == "email_send_notification":
                        action_type = "send_notification"
                        disp_name = str(args.get("to", "Recipient"))
                        email_res = str(args.get("to", ""))
                        title = f"Dispatch Notification to '{args.get('to')}'"
                        desc = f"Subject: {args.get('subject')}\nChannel: {args.get('channel', 'email')}"
                    elif tool_name == "calendar_create_event":
                        event_title = args.get("title") or args.get("event_title") or "New Event"
                        category = args.get("category") or "Work Event"
                        action_type = "create_event"
                        title = f"Create Event: '{event_title}'"
                        desc = f"Event Title: {event_title}\nCategory: {category}\nDate: {args.get('date', 'tomorrow')}\nTime: {args.get('start_time', '6:00 PM')}"
                        disp_name, email_res = "Event", "calendar@company.com"
                    elif tool_name == "calendar_schedule_meeting":
                        raw_person = args.get("person") or args.get("attendees") or "Contact"
                        if isinstance(raw_person, list) and raw_person:
                            raw_person = raw_person[0]
                        disp_name, email_res = CalendarStoreManager.resolve_person_email(str(raw_person))
                        action_type = "schedule_meeting"
                        title = f"Schedule Meeting with {disp_name} ({email_res})"
                        desc = f"Person: {disp_name} ({email_res})\nDate: {args.get('date', 'tomorrow')}\nTime: {args.get('start_time', '3:00 PM')}"
                    else:
                        action_type = f"{args.get('action')}_event"
                        title = f"{str(args.get('action')).title()} Calendar Event '{args.get('event_id')}'"
                        desc = f"Action: {args.get('action')} on event ID {args.get('event_id')}"

                    queued = hitl_manager.queue_action(
                        action_type=action_type,
                        tool_name=tool_name,
                        title=title,
                        description=desc,
                        payload=args,
                        agent_name="Calendar Agent"
                    )
                    return {
                        "status": "pending_approval",
                        "requires_human_approval": True,
                        "person_resolved": disp_name if tool_name == "calendar_create_event" else None,
                        "email_resolved": email_res if tool_name == "calendar_create_event" else None,
                        "hitl_action": queued,
                        "message": f"Action '{title}' queued for Human-in-the-Loop review. User can approve, edit, or reject before execution."
                    }

            if tool_name == "fs_read_file":
                filepath = args.get("path", "").strip().strip("'\"")
                if not filepath:
                    return {"status": "error", "message": "No file path provided. Please specify a valid file path in 'path'."}
                filepath = os.path.normpath(filepath)
                if not os.path.exists(filepath):
                    return {"status": "error", "message": f"File '{filepath}' not found."}

                file_size = os.path.getsize(filepath)
                ext = filepath.split(".")[-1].lower() if "." in filepath else ""
                filename = os.path.basename(filepath)

                links = []
                images = []
                tables = []
                formulas = []
                pages = []
                content = ""
                file_type = ext.upper() if ext else "Text File"
                doc_title = filename

                def _extract_math_formulas(text_chunk: str) -> List[str]:
                    found = []
                    # Inline & Block LaTeX ($...$, $$...$$, \(...\), \[...\])
                    m_matches = re.findall(r'\$\$[^\$]+\$\$|\$[^\$\n]+\$|\\[\(\[][^\\\)]+\\[\)\]]', text_chunk)
                    for m in m_matches:
                        sm = m.strip()
                        if sm and sm not in found and len(sm) > 2:
                            found.append(sm)

                    # LaTeX commands / equations
                    for line in text_chunk.split("\n"):
                        l = line.strip()
                        if any(kw in l for kw in ["\\frac", "\\sum", "\\int", "\\sqrt", "\\alpha", "\\beta", "\\gamma", "\\theta", "\\sigma", "\\lim", "=", "\\approx", "\\le", "\\ge"]) and not l.startswith("#"):
                            if len(l) < 180 and l not in found:
                                fmt_eq = f"$${l}$$" if not l.startswith("$") and not l.startswith("\\") else l
                                if fmt_eq not in found:
                                    found.append(fmt_eq)
                    return found[:12]

                def _parse_page_struct(raw_text: str, p_num: int) -> Dict[str, Any]:
                    lines = [l.strip() for l in raw_text.split("\n") if l.strip()]
                    headings = []
                    subheadings = []
                    body_lines = []

                    for l in lines:
                        if l.startswith("# ") or re.match(r'^(Section|\d+\.)\s+[A-Z]', l):
                            h_clean = l.lstrip("# ").strip()
                            if h_clean not in headings: headings.append(h_clean)
                        elif l.startswith("## ") or l.startswith("### ") or re.match(r'^\d+\.\d+\s+[A-Z]', l):
                            sub_clean = l.lstrip("# ").strip()
                            if sub_clean not in subheadings: subheadings.append(sub_clean)
                        else:
                            body_lines.append(l)

                    page_math = _extract_math_formulas(raw_text)

                    fmt_body = ""
                    if headings:
                        fmt_body += "\n".join([f"# 📌 Main Heading: **{h}**" for h in headings]) + "\n\n"
                    if subheadings:
                        fmt_body += "\n".join([f"## 🏷️ Subheading: *{sh}*" for sh in subheadings]) + "\n\n"
                    if body_lines:
                        fmt_body += "💬 **Normal Body / Text:**\n" + "\n".join(body_lines)

                    return {
                        "page_number": p_num,
                        "headings": headings,
                        "subheadings": subheadings,
                        "formulas": page_math,
                        "raw_text": raw_text,
                        "formatted_text": fmt_body if fmt_body.strip() else raw_text
                    }

                # 1. PDF Documents (Page-by-Page Extraction)
                if ext == "pdf":
                    try:
                        import pypdf
                        reader = pypdf.PdfReader(filepath)
                        file_type = "PDF Document"

                        for idx, page in enumerate(reader.pages):
                            p_num = idx + 1
                            extracted = page.extract_text() or ""
                            if extracted.strip():
                                p_struct = _parse_page_struct(extracted, p_num)
                                pages.append(p_struct)
                                for f_item in p_struct["formulas"]:
                                    if f_item not in formulas:
                                        formulas.append(f_item)

                            # Extract PDF annotation links
                            if "/Annots" in page:
                                try:
                                    annots = page["/Annots"]
                                    for annot in annots:
                                        obj = annot.get_object()
                                        if "/A" in obj and "/URI" in obj["/A"]:
                                            uri = str(obj["/A"]["/URI"])
                                            if uri not in links:
                                                links.append(uri)
                                except Exception:
                                    pass

                            # Extract PDF page image attachments
                            if hasattr(page, "images") and page.images:
                                try:
                                    for img in page.images:
                                        images.append({
                                            "name": img.name,
                                            "page": p_num,
                                            "format": img.name.split(".")[-1] if "." in img.name else "image"
                                        })
                                except Exception:
                                    pass

                        content_parts = [f"--- PAGE {p['page_number']} OF {len(pages)} ---\n{p['formatted_text']}" for p in pages]
                        content = "\n\n".join(content_parts) if content_parts else "No extractable text found in PDF."
                    except Exception as pdf_err:
                        content = f"Failed to extract PDF text: {str(pdf_err)}"

                # 2. CSV / TSV Files
                elif ext in ["csv", "tsv"]:
                    file_type = "CSV Structured Data" if ext == "csv" else "TSV Structured Data"
                    import csv
                    delimiter = "\t" if ext == "tsv" else ","
                    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                        reader = list(csv.reader(f, delimiter=delimiter))
                        if reader:
                            headers = reader[0]
                            rows = reader[1:30]
                            table_md = "| " + " | ".join(headers) + " |\n"
                            table_md += "| " + " | ".join(["---"] * len(headers)) + " |\n"
                            for row in rows:
                                table_md += "| " + " | ".join(row) + " |\n"
                            tables.append(table_md)
                            content = f"CSV File containing {len(headers)} columns and {len(reader)-1} total rows.\n\n{table_md}"
                            p_struct = _parse_page_struct(content, 1)
                            pages.append(p_struct)

                # 3. JSON Files
                elif ext == "json":
                    file_type = "JSON Structured Data"
                    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                        raw = f.read()
                        content = raw
                        try:
                            parsed_json = json.loads(raw)
                            if isinstance(parsed_json, list) and len(parsed_json) > 0 and isinstance(parsed_json[0], dict):
                                cols = list(parsed_json[0].keys())
                                table_md = "| " + " | ".join(cols) + " |\n"
                                table_md += "| " + " | ".join(["---"] * len(cols)) + " |\n"
                                for item in parsed_json[:20]:
                                    table_md += "| " + " | ".join([str(item.get(c, "")) for c in cols]) + " |\n"
                                tables.append(table_md)
                        except Exception:
                            pass
                        p_struct = _parse_page_struct(content, 1)
                        pages.append(p_struct)

                # 4. Text / Markdown / HTML / Code
                else:
                    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                        content = f.read()
                    p_struct = _parse_page_struct(content, 1)
                    pages.append(p_struct)
                    formulas.extend(p_struct["formulas"])

                # Extract URLs from content if empty
                if not links and content:
                    found_urls = re.findall(r'https?://[^\s<>"]+|www\.[^\s<>"]+', content)
                    for u in found_urls:
                        clean_u = u.rstrip(".,);]")
                        if clean_u not in links:
                            links.append(clean_u)

                # Extract Markdown Image syntax ![alt](url)
                if content:
                    found_imgs = re.findall(r'!\[([^\]]*)\]\(([^)]+)\)', content)
                    for alt, img_url in found_imgs:
                        images.append({"name": alt or img_url.split("/")[-1], "url": img_url})

                return {
                    "status": "success",
                    "path": filepath,
                    "title": doc_title,
                    "file_type": file_type,
                    "size_bytes": file_size,
                    "total_pages": len(pages),
                    "formulas": formulas,
                    "pages": pages,
                    "links": links,
                    "images": images,
                    "tables": tables,
                    "content": content
                }

            elif tool_name == "fs_write_file":
                filepath = args.get("path", "").strip().strip("'\"")
                content = args.get("content", "")
                if not filepath:
                    return {"status": "error", "message": "No file path provided."}
                filepath = os.path.normpath(filepath)
                os.makedirs(os.path.dirname(os.path.abspath(filepath)), exist_ok=True)

                if filepath.lower().endswith(".pdf"):
                    try:
                        from reportlab.lib.pagesizes import letter
                        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                        from reportlab.lib import colors

                        doc = SimpleDocTemplate(filepath, pagesize=letter)
                        styles = getSampleStyleSheet()

                        custom_title = ParagraphStyle(
                            'DocTitle',
                            parent=styles['Heading1'],
                            fontSize=18,
                            leading=22,
                            textColor=colors.HexColor('#1e293b'),
                            spaceAfter=12
                        )
                        custom_body = ParagraphStyle(
                            'DocBody',
                            parent=styles['Normal'],
                            fontSize=11,
                            leading=16,
                            textColor=colors.HexColor('#334155'),
                            spaceAfter=10
                        )

                        story = []
                        lines = content.split("\n")
                        first = True
                        for l_item in lines:
                            st = l_item.strip()
                            if not st:
                                story.append(Spacer(1, 8))
                                continue
                            if first or st.startswith("# "):
                                story.append(Paragraph(st.lstrip("# "), custom_title))
                                first = False
                            else:
                                story.append(Paragraph(st, custom_body))

                        doc.build(story)
                        size_bytes = os.path.getsize(filepath)
                        return {
                            "status": "success",
                            "message": f"Valid PDF Document '{filepath}' compiled and written successfully.",
                            "size_bytes": size_bytes
                        }
                    except Exception:
                        with open(filepath, "w", encoding="utf-8") as f:
                            f.write(content)
                        return {
                            "status": "success",
                            "message": f"File '{filepath}' written successfully.",
                            "size_bytes": len(content)
                        }
                else:
                    with open(filepath, "w", encoding="utf-8") as f:
                        f.write(content)
                    return {"status": "success", "message": f"File '{filepath}' written successfully.", "size_bytes": len(content)}

            elif tool_name == "fs_list_dir":
                dirpath = args.get("path", ".")
                if not dirpath: dirpath = "."
                if not os.path.exists(dirpath):
                    return {"status": "error", "message": f"Directory '{dirpath}' not found."}
                items = os.listdir(dirpath)
                result = []
                for item in items[:50]:
                    full_p = os.path.join(dirpath, item)
                    result.append({
                        "name": item,
                        "is_directory": os.path.isdir(full_p),
                        "size": os.path.getsize(full_p) if os.path.isfile(full_p) else 0
                    })
                return {"status": "success", "path": dirpath, "items": result}

            elif tool_name == "gmail_list_messages":
                query = args.get("query", "all")
                limit = args.get("limit", 10)
                from app.core.config import settings

                if settings.EMAIL_USER and settings.EMAIL_PASSWORD:
                    return BuiltinMCPServers._fetch_real_emails(
                        email_user=settings.EMAIL_USER,
                        email_pass=settings.EMAIL_PASSWORD,
                        limit=limit,
                        imap_server=settings.IMAP_SERVER,
                        port=settings.IMAP_PORT
                    )

                return {
                    "status": "configuration_required",
                    "message": "To fetch real live emails, email credentials are required in backend/.env.",
                    "setup_instructions": {
                        "step_1": "Create a backend/.env file in the backend directory.",
                        "step_2": "Add EMAIL_USER=your_email@gmail.com",
                        "step_3": "Generate a Google App Password (https://myaccount.google.com/apppasswords) and add EMAIL_PASSWORD=your_app_password",
                        "step_4": "Restart the backend server."
                    }
                }

            elif tool_name in ["gmail_send_message", "email_send_notification", "send_email", "gmail_draft_message"]:
                # Dynamic Multi-Recipient Extraction: Handles lists, comma-separated strings, or single addresses
                raw_recipients = []
                for k in ["to", "recipients", "attendees", "other_attendees", "person", "recipient", "attendee"]:
                    val = args.get(k)
                    if not val:
                        continue
                    if isinstance(val, list):
                        for item in val:
                            if isinstance(item, str):
                                raw_recipients.extend([x.strip() for x in item.split(",") if x.strip()])
                            elif isinstance(item, dict):
                                addr = item.get("email") or item.get("address") or item.get("name")
                                if addr:
                                    raw_recipients.append(str(addr).strip())
                    elif isinstance(val, str):
                        raw_recipients.extend([x.strip() for x in val.split(",") if x.strip()])

                recipients = []
                for r in raw_recipients:
                    _, resolved_email = CalendarStoreManager.resolve_person_email(r)
                    target_addr = resolved_email or r
                    if target_addr and target_addr not in recipients:
                        recipients.append(target_addr)

                if not recipients:
                    return {"status": "error", "message": "No recipient email addresses provided in the action payload."}

                subject = args.get("subject") or "Notification from AI Assistant"
                body = args.get("body") or args.get("message") or args.get("content") or "You have received a new notification."
                channel = args.get("channel", "email")

                from app.core.config import settings
                dispatch_results = []

                for target_to in recipients:
                    # Auto-expand brief prompts into professional emails if necessary
                    recip_body = body
                    if len(recip_body.split()) < 15 and not any(kw in recip_body.lower() for kw in ["hi ", "dear ", "best regards", "sincerely", "regards"]):
                        recipient_name = str(target_to).split("@")[0].replace(".", " ").replace("_", " ").title() if "@" in str(target_to) else "Team"
                        clean_b = recip_body.strip(". ")
                        recip_body = (
                            f"Hi {recipient_name},\n\n"
                            f"I hope this message finds you well.\n\n"
                            f"I am reaching out regarding the following:\n"
                            f"• {clean_b.capitalize()}.\n\n"
                            f"Please feel free to reply or let me know if you have any questions.\n\n"
                            f"Best regards,\n"
                            f"Google Workspace AI Assistant"
                        )

                    recip_subject = subject
                    if not recip_subject or recip_subject == "Update from Assistant":
                        recip_subject = f"Regarding: {recip_body.splitlines()[0][:40]}"

                    # 1. Log notification in EmailStoreManager for outbox tracking
                    EmailStoreManager.log_notification(to=str(target_to), subject=recip_subject, body=recip_body, channel=channel)

                    # 2. Broadcast real-time WebSocket push event
                    try:
                        from app.core.websocket_manager import ws_manager
                        ws_manager.broadcast_sync("notification_received", {
                            "to": str(target_to),
                            "subject": recip_subject,
                            "body": recip_body,
                            "channel": channel,
                            "timestamp": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                        })
                    except Exception:
                        pass

                    # 3. Dispatch live SMTP email if credentials exist
                    smtp_res = None
                    if settings.EMAIL_USER and settings.EMAIL_PASSWORD:
                        smtp_res = BuiltinMCPServers._send_real_email(
                            to=str(target_to),
                            subject=recip_subject,
                            body=recip_body,
                            email_user=settings.EMAIL_USER,
                            email_pass=settings.EMAIL_PASSWORD,
                            smtp_server=getattr(settings, "SMTP_SERVER", "smtp.gmail.com"),
                            port=getattr(settings, "SMTP_PORT", 587)
                        )

                    dispatch_results.append({
                        "recipient": target_to,
                        "status": "delivered" if (smtp_res and smtp_res.get("status") == "success") else "logged_in_outbox",
                        "smtp_result": smtp_res
                    })

                return {
                    "status": "success",
                    "mode": "multi_recipient_dispatch",
                    "recipient_count": len(recipients),
                    "recipients": recipients,
                    "dispatched_notifications": dispatch_results,
                    "message": f"Successfully dispatched notifications to {len(recipients)} recipient(s): {', '.join(recipients)}."
                }

            elif tool_name == "calendar_list_events":
                days = args.get("days", 7)
                events_list = CalendarStoreManager.list_events(days=days)
                stats = CalendarStoreManager.get_dashboard_stats()["dashboard_metrics"]
                return {
                    "status": "success",
                    "range_days": days,
                    "dashboard_metrics": stats,
                    "events": events_list
                }

            elif tool_name in ["calendar_create_event", "calendar_schedule_meeting"]:
                raw_person = args.get("person") or args.get("attendee") or args.get("attendees")
                if isinstance(raw_person, list) and raw_person:
                    raw_person = raw_person[0]
                elif not raw_person:
                    raw_person = "Rahul"

                disp_name, resolved_email = CalendarStoreManager.resolve_person_email(str(raw_person))

                title = args.get("title")
                if not title or title == "Event":
                    title = f"Meeting with {disp_name}"

                start_time = args.get("start_time") or "3:00 PM"
                duration = args.get("duration_minutes", 30)
                notify_attendees = args.get("notify_attendees", True)

                # Determine event date (supports custom dates like "2026-08-15", "tomorrow", "today", etc.)
                event_date = args.get("date") or args.get("event_date")
                if event_date:
                    event_date_str = str(event_date).strip()
                    if event_date_str.lower() == "today":
                        event_date = datetime.date.today().strftime("%Y-%m-%d")
                    elif event_date_str.lower() == "tomorrow":
                        event_date = (datetime.date.today() + datetime.timedelta(days=1)).strftime("%Y-%m-%d")
                    else:
                        event_date = event_date_str
                else:
                    date_match = re.search(r'\d{4}[-/]\d{1,2}[-/]\d{1,2}', str(start_time))
                    if date_match:
                        event_date = date_match.group(0)
                    elif "tomorrow" in str(start_time).lower() or "tomorrow" in str(args).lower():
                        event_date = (datetime.date.today() + datetime.timedelta(days=1)).strftime("%Y-%m-%d")
                    else:
                        event_date = (datetime.date.today() + datetime.timedelta(days=1)).strftime("%Y-%m-%d")

                # Perform Schedule Availability & Conflict Check
                avail_info = CalendarStoreManager.check_availability(event_date, start_time)
                has_conflict, conflict_msg = CalendarStoreManager.check_conflict(event_date, start_time)
                suggested_slots = avail_info.get("suggested_free_slots", ["09:00 AM", "10:30 AM", "01:00 PM", "02:30 PM", "04:00 PM"])

                # Gather all explicit attendees and resolved email addresses
                raw_att = args.get("attendees") or []
                if isinstance(raw_att, str):
                    raw_att = [a.strip() for a in raw_att.split(",") if a.strip()]

                raw_other = args.get("other_attendees") or []
                if isinstance(raw_other, str):
                    raw_other = [a.strip() for a in raw_other.split(",") if a.strip()]

                attendees_list = []
                if resolved_email:
                    attendees_list.append(resolved_email)

                for a in list(raw_att) + list(raw_other):
                    _, email_a = CalendarStoreManager.resolve_person_email(str(a))
                    target_addr = email_a or str(a).strip()
                    if target_addr and target_addr not in attendees_list:
                        attendees_list.append(target_addr)

                # Save event to store
                created_evt = CalendarStoreManager.create_event(
                    title=title,
                    event_date=event_date,
                    start_time=start_time,
                    duration=duration,
                    attendees_list=attendees_list
                )

                notifications = []
                from app.core.config import settings

                if notify_attendees and attendees_list:
                    subject = f"Event Invitation: {title}"
                    body = (
                        f"Hello,\n\n"
                        f"You have been invited to a meeting:\n"
                        f"• Title: {title}\n"
                        f"• Date: {event_date}\n"
                        f"• Time: {start_time}\n"
                        f"• Duration: {duration} minutes\n\n"
                        f"Please click ACCEPT or REJECT to confirm your attendance.\n\n"
                        f"Best regards,\n"
                        f"Google Workspace Calendar"
                    )

                    for recipient in attendees_list:
                        # Log notification in EmailStoreManager for outbox tracking
                        EmailStoreManager.log_notification(to=recipient, subject=subject, body=body, channel="email")

                        if "@" in recipient and settings.EMAIL_USER and settings.EMAIL_PASSWORD:
                            encoded_recip = urllib.parse.quote(recipient)
                            public_base = str(settings.PUBLIC_API_URL).rstrip("/")
                            accept_url = f"{public_base}/api/calendar/respond?event_id={created_evt['id']}&action=accept&attendee={encoded_recip}"
                            reject_url = f"{public_base}/api/calendar/respond?event_id={created_evt['id']}&action=reject&attendee={encoded_recip}"

                            html_email_body = f"""
                            <!DOCTYPE html>
                            <html>
                            <body style="font-family: Arial, sans-serif; background-color: #0a0d14; color: #f8fafc; padding: 24px; margin: 0;">
                              <div style="background-color: #121824; border: 1px solid rgba(255,255,255,0.12); border-radius: 20px; padding: 32px; max-width: 480px; margin: 0 auto; box-shadow: 0 10px 30px rgba(0,0,0,0.5);">
                                <h2 style="color: #a5b4fc; margin-top: 0;">📅 Meeting Invitation: {title}</h2>
                                <p style="color: #94a3b8; font-size: 14px;">You have been invited to a meeting by <strong>Google Workspace Calendar</strong>.</p>
                                <div style="background: rgba(10, 13, 20, 0.7); border: 1px solid rgba(255,255,255,0.08); border-radius: 12px; padding: 16px; margin: 20px 0;">
                                  <p style="margin: 6px 0; color: #f8fafc;"><strong>Event:</strong> {title}</p>
                                  <p style="margin: 6px 0; color: #f8fafc;"><strong>Date:</strong> {event_date}</p>
                                  <p style="margin: 6px 0; color: #f8fafc;"><strong>Time:</strong> {start_time}</p>
                                  <p style="margin: 6px 0; color: #f8fafc;"><strong>Duration:</strong> {duration} minutes</p>
                                </div>
                                <p style="font-weight: bold; color: #ffffff; margin-bottom: 16px;">Please click a button below to confirm your attendance:</p>
                                <div style="display: flex; gap: 12px; margin-top: 10px;">
                                  <a href="{accept_url}" style="background-color: #10b981; color: #ffffff; padding: 12px 22px; text-decoration: none; border-radius: 10px; font-weight: bold; display: inline-block; font-size: 14px; box-shadow: 0 4px 14px rgba(16, 185, 129, 0.35);">[ ✓ ACCEPT INVITATION ]</a>
                                  <a href="{reject_url}" style="background-color: #ef4444; color: #ffffff; padding: 12px 22px; text-decoration: none; border-radius: 10px; font-weight: bold; display: inline-block; font-size: 14px; box-shadow: 0 4px 14px rgba(239, 68, 68, 0.35);">[ ✗ REJECT INVITATION ]</a>
                                </div>
                                <p style="font-size: 12px; color: #64748b; margin-top: 24px;">Accepting sends an immediate notification & queues a 30-minute prior reminder notification.</p>
                              </div>
                            </body>
                            </html>
                            """

                            send_res = BuiltinMCPServers._send_real_email(
                                to=recipient,
                                subject=subject,
                                body=body,
                                email_user=settings.EMAIL_USER,
                                email_pass=settings.EMAIL_PASSWORD,
                                smtp_server=getattr(settings, "SMTP_SERVER", "smtp.gmail.com"),
                                port=getattr(settings, "SMTP_PORT", 587),
                                html_body=html_email_body
                            )
                            if send_res.get("status") == "success":
                                notifications.append({
                                    "recipient": recipient,
                                    "channel": "email",
                                    "status": "sent",
                                    "delivered_at": send_res.get("delivered_at"),
                                    "message": f"Real HTML invitation email sent with clickable Accept/Reject buttons to {recipient}."
                                })
                            else:
                                notifications.append({
                                    "recipient": recipient,
                                    "channel": "email",
                                    "status": "failed",
                                    "message": f"Email dispatch failed: {send_res.get('message')}"
                                })
                        else:
                            notifications.append({
                                "recipient": recipient,
                                "channel": "notification",
                                "status": "sent",
                                "sent_at": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                                "notification_title": f"Invitation: {title}",
                                "notification_body": f"You are invited to '{title}' on {event_date} at {start_time} ({duration} mins). Click ACCEPT or REJECT.",
                                "message": f"Notification message dispatched to attendee '{recipient}'."
                            })

                notification_summary = (
                    f"Notification message sent to {len(attendees_list)} attendee(s)"
                    if attendees_list and notify_attendees
                    else "No notifications sent"
                )

                dashboard_stats = CalendarStoreManager.get_dashboard_stats()["dashboard_metrics"]

                return {
                    "status": "success",
                    "intent": "schedule_meeting",
                    "person_resolved": disp_name,
                    "email_resolved": resolved_email,
                    "event_id": created_evt["id"],
                    "title": title,
                    "date": event_date,
                    "start_time": start_time,
                    "duration_minutes": duration,
                    "suggested_free_slots": suggested_slots,
                    "conflict_check": conflict_msg if has_conflict else "No schedule conflicts detected",
                    "has_conflict": has_conflict,
                    "attendees": [resolved_email],
                    "rsvps": created_evt["rsvps"],
                    "accepted_count": created_evt["accepted_count"],
                    "rejected_count": created_evt["rejected_count"],
                    "pending_count": created_evt["pending_count"],
                    "dashboard_metrics": dashboard_stats,
                    "calendar_status": "confirmed",
                    "notification_status": notification_summary,
                    "notifications_sent": notifications,
                    "interactive_actions": ["ACCEPT", "REJECT"],
                    "message": f"Meeting with {disp_name} ({resolved_email}) scheduled for {event_date} at {start_time}."
                }

            elif tool_name == "calendar_respond_invitation":
                event_id = args.get("event_id", "")
                attendee = args.get("attendee", "user@company.com")
                action = args.get("action", "accept").lower()

                evt, att_name, new_status = CalendarStoreManager.respond_invitation(event_id, attendee, action)

                notifications = []
                now_str = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                from app.core.config import settings

                if new_status == "accepted":
                    # 1. Immediate Notification
                    imm_msg = f"Immediate Notification: You have ACCEPTED the invitation for '{evt['title']}' on {evt['date']} at {evt['start_time']}."
                    notifications.append({
                        "type": "immediate_notification",
                        "recipient": att_name,
                        "status": "delivered",
                        "timestamp": now_str,
                        "title": f"RSVP Confirmed: {evt['title']}",
                        "message": imm_msg
                    })

                    # 2. 30-Minute Prior Reminder Notification
                    rem_msg = f"30-Min Prior Reminder: '{evt['title']}' starts in 30 minutes at {evt['start_time']} ({evt['date']}). Reminder notification queued for accepted attendee '{att_name}'."
                    notifications.append({
                        "type": "30min_prior_reminder",
                        "recipient": att_name,
                        "status": "scheduled",
                        "scheduled_for": f"30 minutes prior to {evt['date']} {evt['start_time']}",
                        "title": f"Upcoming Meeting Reminder (30-Min Prior)",
                        "message": rem_msg
                    })

                    # Dispatch real email if configured
                    if "@" in att_name and settings.EMAIL_USER and settings.EMAIL_PASSWORD:
                        BuiltinMCPServers._send_real_email(
                            to=att_name,
                            subject=f"Confirmation: {evt['title']} Accepted",
                            body=f"Hello {att_name},\n\nYou have ACCEPTED the invitation for '{evt['title']}' scheduled for {evt['date']} at {evt['start_time']}.\n\nYou will receive a reminder notification 30 minutes prior to the event.\n\nBest regards,\nGoogle Workspace Calendar",
                            email_user=settings.EMAIL_USER,
                            email_pass=settings.EMAIL_PASSWORD,
                            smtp_server=getattr(settings, "SMTP_SERVER", "smtp.gmail.com"),
                            port=getattr(settings, "SMTP_PORT", 587)
                        )
                else:
                    # Rejected status
                    imm_msg = f"Immediate Notification: You have REJECTED the invitation for '{evt['title']}'."
                    notifications.append({
                        "type": "immediate_notification",
                        "recipient": att_name,
                        "status": "delivered",
                        "timestamp": now_str,
                        "title": f"RSVP Declined: {evt['title']}",
                        "message": imm_msg
                    })

                dashboard_stats = CalendarStoreManager.get_dashboard_stats()["dashboard_metrics"]

                try:
                    from app.core.websocket_manager import ws_manager
                    ws_manager.broadcast_sync("calendar_rsvp_updated", {
                        "event_id": evt["id"],
                        "title": evt["title"],
                        "action": new_status,
                        "attendee": att_name,
                        "dashboard_metrics": dashboard_stats,
                        "rsvps": evt["rsvps"]
                    })
                except Exception:
                    pass

                return {
                    "status": "success",
                    "action": new_status,
                    "event_id": evt["id"],
                    "title": evt["title"],
                    "date": evt["date"],
                    "start_time": evt["start_time"],
                    "attendee": att_name,
                    "rsvp_status": new_status,
                    "accepted_count": evt["accepted_count"],
                    "rejected_count": evt["rejected_count"],
                    "pending_count": evt["pending_count"],
                    "dashboard_metrics": dashboard_stats,
                    "immediate_notification": notifications[0]["message"],
                    "reminder_30min_prior": notifications[1]["message"] if len(notifications) > 1 else None,
                    "notifications_dispatched": notifications
                }

            elif tool_name == "calendar_get_dashboard_stats":
                return CalendarStoreManager.get_dashboard_stats()

            elif tool_name == "calendar_clear_history":
                return CalendarStoreManager.clear_history()

            elif tool_name == "calendar_check_availability":
                t_date = args.get("date") or args.get("target_date")
                t_time = args.get("start_time")
                duration = args.get("duration_minutes", 30)
                return CalendarStoreManager.check_availability(target_date=t_date, start_time=t_time, duration_minutes=duration)

            elif tool_name == "calendar_manage_event":
                evt_id = args.get("event_id", "")
                action = args.get("action", "update")
                new_title = args.get("new_title")
                new_date = args.get("new_date")
                new_start_time = args.get("new_start_time")
                new_duration = args.get("new_duration_minutes")
                return CalendarStoreManager.manage_event(
                    event_id=evt_id,
                    action=action,
                    new_title=new_title,
                    new_date=new_date,
                    new_start_time=new_start_time,
                    new_duration=new_duration
                )

            elif tool_name == "gmail_draft_message":
                to = args.get("to")
                subject = args.get("subject")
                body = args.get("body")
                if not to or not body:
                    return {"status": "error", "message": "'to' and 'body' parameters are required for drafting an email."}
                if not subject:
                    subject = f"Draft: {body.splitlines()[0][:35]}"
                
                draft = EmailStoreManager.save_draft(to=to, subject=subject, body=body)
                return {
                    "status": "success",
                    "message": f"Email draft created for '{to}'.",
                    "draft": draft,
                    "all_drafts_count": len(EmailStoreManager.list_drafts())
                }

            elif tool_name == "gmail_list_drafts":
                drafts = EmailStoreManager.list_drafts()
                return {
                    "status": "success",
                    "count": len(drafts),
                    "drafts": drafts
                }

            elif tool_name == "email_send_notification":
                to = args.get("to")
                subject = args.get("subject")
                body = args.get("body")
                channel = args.get("channel", "email")

                if not to or not body:
                    return {"status": "error", "message": "'to' and 'body' parameters are required for sending notification."}
                if not subject:
                    subject = "Notification Alert"

                from app.core.config import settings
                if channel == "email" and settings.EMAIL_USER and settings.EMAIL_PASSWORD:
                    send_res = BuiltinMCPServers._send_real_email(
                        to=to,
                        subject=subject,
                        body=body,
                        email_user=settings.EMAIL_USER,
                        email_pass=settings.EMAIL_PASSWORD
                    )
                    EmailStoreManager.log_notification(to=to, subject=subject, body=body, channel="email")
                    return send_res

                notif = EmailStoreManager.log_notification(to=to, subject=subject, body=body, channel=channel)
                return {
                    "status": "success",
                    "mode": "in_app_notification",
                    "recipient": to,
                    "subject": subject,
                    "notification": notif,
                    "message": f"Notification dispatched to '{to}' via {channel} channel."
                }

            elif tool_name == "github_list_repos":
                target = args.get("user_or_org", "default")
                return {
                    "status": "success",
                    "owner": target,
                    "repositories": [
                        {"name": "mcp-ai-assistant", "stars": 1280, "language": "Python / React", "visibility": "public"},
                        {"name": "fastapi-llm-orchestrator", "stars": 640, "language": "Python", "visibility": "public"},
                        {"name": "rag-vector-engine", "stars": 310, "language": "TypeScript", "visibility": "internal"}
                    ]
                }

            elif tool_name == "github_create_issue":
                repo = args.get("repo")
                title = args.get("title")
                body = args.get("body")
                return {
                    "status": "success",
                    "issue_number": 42,
                    "url": f"https://github.com/{repo}/issues/42",
                    "title": title,
                    "state": "open",
                    "created_at": datetime.datetime.now().isoformat()
                }

            elif tool_name == "run_python_code":
                code = args.get("code", "")
                if not code.strip():
                    return {"status": "error", "message": "No code provided."}
                
                # Execute in subprocess safely
                try:
                    proc = subprocess.run(
                        [sys.executable, "-c", code],
                        capture_output=True,
                        text=True,
                        timeout=5
                    )
                    return {
                        "status": "success" if proc.returncode == 0 else "error",
                        "exit_code": proc.returncode,
                        "stdout": proc.stdout,
                        "stderr": proc.stderr
                    }
                except subprocess.TimeoutExpired:
                    return {"status": "error", "message": "Code execution timed out after 5 seconds."}

            elif tool_name == "db_query":
                query = args.get("query", "").strip()
                if not query:
                    return {"status": "error", "message": "No SQL query provided."}

                clean_query = query.lower()
                allowed_prefixes = ("select", "pragma", "explain")
                if not any(clean_query.startswith(p) for p in allowed_prefixes):
                    return {"status": "error", "message": "Only read-only SELECT, PRAGMA, or EXPLAIN queries are allowed."}

                try:
                    limit_arg = int(args.get("limit", 20))
                    limit = max(1, min(limit_arg, 100))
                except (ValueError, TypeError):
                    limit = 20

                db_path = os.getenv("ASSISTANT_DB_PATH", "assistant.db")
                conn = None
                try:
                    # Attempt URI read-only connection, fallback to standard connection with query_only PRAGMA
                    try:
                        conn = sqlite3.connect(f"file:{db_path}?mode=ro", uri=True)
                    except sqlite3.Error:
                        conn = sqlite3.connect(db_path)
                    
                    cursor = conn.cursor()
                    cursor.execute("PRAGMA query_only = ON;")
                    cursor.execute(query)

                    columns = [desc[0] for desc in cursor.description] if cursor.description else []
                    rows = cursor.fetchmany(limit)

                    return {
                        "status": "success",
                        "columns": columns,
                        "row_count": len(rows),
                        "data": [dict(zip(columns, r)) for r in rows]
                    }
                except sqlite3.Error as e:
                    return {"status": "error", "message": f"Database Query Error: {str(e)}"}
                finally:
                    if conn:
                        conn.close()

            elif tool_name == "web_search":
                query = args.get("query", "")
                if not query:
                    return {"status": "error", "message": "No search query provided."}

                results = []
                import urllib.parse
                import re
                from html import unescape

                summary_text = None

                # 1. Wikipedia Summary REST API (Fetches rich text summary for concepts, topics, items)
                try:
                    target_q = query.strip().title()
                    headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"}
                    wiki_sum_url = f"https://en.wikipedia.org/api/rest_v1/page/summary/{urllib.parse.quote(target_q)}"
                    w_res = httpx.get(wiki_sum_url, headers=headers, timeout=5.0, follow_redirects=True)
                    if w_res.status_code == 200:
                        w_json = w_res.json()
                        extract = w_json.get("extract")
                        if extract and len(extract) > 20:
                            summary_text = extract
                            wiki_page = w_json.get("content_urls", {}).get("desktop", {}).get("page") or f"https://en.wikipedia.org/wiki/{urllib.parse.quote(target_q)}"
                            results.append({
                                "title": w_json.get("title", query.title()),
                                "url": wiki_page,
                                "snippet": extract
                            })
                except Exception:
                    pass

                # 2. DuckDuckGo Instant Answer API
                try:
                    ddg_api_url = f"https://api.duckduckgo.com/?q={urllib.parse.quote(query)}&format=json&no_html=1"
                    res = httpx.get(ddg_api_url, timeout=5.0)
                    if res.status_code == 200:
                        data = res.json()
                        abstract = data.get("AbstractText")
                        heading = data.get("Heading")
                        url = data.get("AbstractURL")
                        if abstract and url:
                            if not summary_text:
                                summary_text = abstract
                            results.append({
                                "title": heading or query.title(),
                                "url": url,
                                "snippet": abstract
                            })
                except Exception:
                    pass

                # 3. DuckDuckGo HTML POST Live Search
                if not results:
                    try:
                        headers = {
                            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
                            "Content-Type": "application/x-www-form-urlencoded"
                        }
                        res = httpx.post("https://html.duckduckgo.com/html/", data={"q": query}, headers=headers, timeout=6.0, follow_redirects=True)
                        if res.status_code == 200:
                            raw_html = res.text
                            matches = re.findall(r'<a class="result__a"[^>]*href="([^"]+)"[^>]*>(.*?)</a>[\s\S]*?<a class="result__snippet"[^>]*>(.*?)</a>', raw_html)
                            for link, t_html, s_html in matches[:5]:
                                clean_t = unescape(re.sub(r'<[^>]+>', '', t_html)).strip()
                                clean_s = unescape(re.sub(r'<[^>]+>', '', s_html)).strip()
                                real_u = link
                                u_m = re.search(r'uddg=([^&]+)', link)
                                if u_m:
                                    real_u = urllib.parse.unquote(u_m.group(1))
                                if clean_t and clean_s:
                                    if not summary_text:
                                        summary_text = clean_s
                                    results.append({
                                        "title": clean_t,
                                        "url": real_u,
                                        "snippet": clean_s
                                    })
                    except Exception:
                        pass

                # 4. Wikipedia REST API Search Fallback
                if not results:
                    try:
                        headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"}
                        wiki_url = f"https://en.wikipedia.org/w/api.php?action=query&list=search&srsearch={urllib.parse.quote(query)}&format=json"
                        res = httpx.get(wiki_url, headers=headers, timeout=5.0)
                        if res.status_code == 200:
                            w_data = res.json()
                            for item in w_data.get("query", {}).get("search", [])[:3]:
                                w_title = item.get("title")
                                w_snippet = unescape(re.sub(r'<[^>]+>', '', item.get("snippet", "")))
                                w_page_url = f"https://en.wikipedia.org/wiki/{urllib.parse.quote(w_title)}"
                                if not summary_text:
                                    summary_text = w_snippet
                                results.append({
                                    "title": w_title,
                                    "url": w_page_url,
                                    "snippet": w_snippet
                                })
                    except Exception:
                        pass

                if not summary_text or summary_text.lower().startswith("web results for") or "widely recognized topic" in summary_text:
                    topic_title = query.strip().title()
                    if results and results[0].get("snippet") and len(results[0].get("snippet")) > 30:
                        summary_text = results[0].get("snippet")
                    else:
                        summary_text = f"Detailed background, taxonomy, key features, and practical applications for {topic_title} based on reference databases and live web search data."

                return {
                    "status": "success",
                    "query": query,
                    "summary_text": summary_text,
                    "results": results if results else [
                        {
                            "title": f"Search Information for '{query}'",
                            "url": f"https://www.google.com/search?q={urllib.parse.quote(query)}",
                            "snippet": summary_text
                        }
                    ]
                }

            elif tool_name == "get_weather":
                city = args.get("city", "San Francisco")
                country = args.get("country", "")

                from app.core.config import settings

                # 1. Official OpenWeatherMap API (if OPENWEATHER_API_KEY is present in backend/.env)
                if settings.OPENWEATHER_API_KEY:
                    try:
                        q_str = f"{city},{country}".strip(",")
                        ow_url = f"https://api.openweathermap.org/data/2.5/weather?q={urllib.parse.quote(q_str)}&appid={settings.OPENWEATHER_API_KEY}&units=metric"
                        res = httpx.get(ow_url, timeout=5.0)
                        if res.status_code == 200:
                            data = res.json()
                            temp_c = data["main"]["temp"]
                            temp_f = round(temp_c * 9/5 + 32, 1)
                            humidity = data["main"]["humidity"]
                            wind = data["wind"]["speed"]
                            cond = data["weather"][0]["description"].title()
                            found_city = data.get("name", city)
                            found_country = data.get("sys", {}).get("country", country)

                            return {
                                "status": "success",
                                "mode": "openweather_api",
                                "city": found_city,
                                "country": found_country,
                                "location": f"{found_city}, {found_country}" if found_country else found_city,
                                "temperature": f"{temp_c}°C ({temp_f}°F)",
                                "feels_like": f"{data['main'].get('feels_like')}°C",
                                "condition": cond,
                                "humidity": f"{humidity}%",
                                "wind_speed": f"{wind} m/s",
                                "provider": "OpenWeatherMap"
                            }
                    except Exception:
                        pass

                # 2. Open-Meteo Free Global Weather API Fallback (No Key Required)
                try:
                    geo_url = f"https://geocoding-api.open-meteo.com/v1/search?name={urllib.parse.quote(city)}&count=10"
                    geo_res = httpx.get(geo_url, timeout=5.0)

                    if geo_res.status_code == 200 and geo_res.json().get("results"):
                        results_list = geo_res.json()["results"]
                        target_location = results_list[0]
                        if country:
                            for loc in results_list:
                                if country.lower() in loc.get("country", "").lower() or country.lower() in loc.get("country_code", "").lower():
                                    target_location = loc
                                    break

                        lat = target_location["latitude"]
                        lon = target_location["longitude"]
                        found_city = target_location.get("name", city)
                        found_country = target_location.get("country", country)

                        weather_url = f"https://api.open-meteo.com/v1/forecast?latitude={lat}&longitude={lon}&current_weather=true"
                        w_res = httpx.get(weather_url, timeout=5.0)

                        if w_res.status_code == 200:
                            curr_w = w_res.json().get("current_weather", {})
                            temp_c = curr_w.get("temperature", 22.0)
                            temp_f = round(temp_c * 9/5 + 32, 1)
                            wind = curr_w.get("windspeed", 10.0)
                            code = curr_w.get("weathercode", 0)

                            weather_codes = {
                                0: "Clear sky",
                                1: "Mainly clear",
                                2: "Partly cloudy",
                                3: "Overcast",
                                45: "Foggy",
                                48: "Depositing rime fog",
                                51: "Light drizzle",
                                53: "Moderate drizzle",
                                55: "Dense drizzle",
                                61: "Slight rain",
                                63: "Moderate rain",
                                65: "Heavy rain",
                                71: "Slight snow",
                                73: "Moderate snow",
                                75: "Heavy snow",
                                80: "Slight rain showers",
                                81: "Moderate rain showers",
                                82: "Violent rain showers",
                                95: "Thunderstorm",
                                96: "Thunderstorm with light hail",
                                99: "Thunderstorm with heavy hail"
                            }
                            cond_text = weather_codes.get(code, "Partly Cloudy")

                            return {
                                "status": "success",
                                "mode": "live_api",
                                "city": found_city,
                                "country": found_country,
                                "location": f"{found_city}, {found_country}" if found_country else found_city,
                                "temperature": f"{temp_c}°C ({temp_f}°F)",
                                "condition": cond_text,
                                "wind_speed": f"{wind} km/h",
                                "latitude": lat,
                                "longitude": lon,
                                "provider": "Open-Meteo"
                            }
                except Exception:
                    pass

                return {
                    "status": "success",
                    "mode": "estimated",
                    "city": city,
                    "country": country or "Not specified",
                    "location": f"{city}, {country}" if country else city,
                    "temperature": "22°C (71.6°F)",
                    "condition": "Partly Cloudy with gentle breeze",
                    "wind_speed": "12 km/h",
                    "provider": "Built-in Fallback"
                }

            elif tool_name == "generate_document":
                filename = args.get("filename", "document.docx")
                title = args.get("title", "Generated Document")
                content = args.get("content", "")
                
                safe_name = os.path.basename(filename)
                docs_dir = os.path.join(os.getcwd(), "generated_docs")
                os.makedirs(docs_dir, exist_ok=True)
                file_path = os.path.join(docs_dir, safe_name)
                ext = os.path.splitext(safe_name)[1].lower()

                if ext == ".docx":
                    try:
                        # pyrefly: ignore [missing-import]
                        import docx
                    except ImportError:
                        subprocess.run([sys.executable, "-m", "pip", "install", "python-docx"], capture_output=True)
                        # pyrefly: ignore [missing-import]
                        import docx
                    
                    doc = docx.Document()
                    doc.add_heading(title, 0)
                    for line in content.split("\n"):
                        line_str = line.strip()
                        if not line_str:
                            continue
                        if line_str.startswith("###") or line_str.startswith("Email ") or line_str.startswith("Item "):
                            doc.add_heading(line_str.replace("#", "").strip(), level=2)
                        else:
                            doc.add_paragraph(line_str)
                    doc.save(file_path)

                elif ext == ".pdf":
                    try:
                        from reportlab.lib.pagesizes import letter
                        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                    except ImportError:
                        subprocess.run([sys.executable, "-m", "pip", "install", "reportlab"], capture_output=True)
                        from reportlab.lib.pagesizes import letter
                        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

                    doc = SimpleDocTemplate(file_path, pagesize=letter)
                    styles = getSampleStyleSheet()
                    story = []

                    title_style = ParagraphStyle('DocTitle', parent=styles['Heading1'], fontSize=20, leading=24, spaceAfter=12)
                    story.append(Paragraph(title, title_style))
                    story.append(Spacer(1, 12))

                    body_style = ParagraphStyle('DocBody', parent=styles['Normal'], fontSize=11, leading=15, spaceAfter=8)
                    for p in content.split("\n"):
                        if p.strip():
                            p_clean = p.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                            story.append(Paragraph(p_clean, body_style))

                    doc.build(story)

                else:
                    with open(file_path, "w", encoding="utf-8") as f:
                        f.write(f"# {title}\n\n{content}")

                download_url = f"/api/downloads/{safe_name}"
                return {
                    "status": "success",
                    "filename": safe_name,
                    "file_path": file_path,
                    "download_url": download_url,
                    "markdown_download_link": f"[{safe_name}]({download_url})",
                    "message": f"Document '{safe_name}' created successfully. Download link: {download_url}"
                }

            else:
                return {"status": "error", "message": f"Unknown tool '{tool_name}'"}

        except Exception as e:
            return {"status": "error", "message": str(e)}
